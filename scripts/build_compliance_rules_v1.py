from __future__ import annotations

import hashlib
import re
import subprocess
import xml.etree.ElementTree as ET
from pathlib import Path
from zipfile import ZipFile

import pandas as pd
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


REPO_ROOT = Path(__file__).resolve().parents[1]
RAW_DIR = REPO_ROOT / "data" / "base-line-0312" / "法规文件_中国"
OUTPUT_DIR = REPO_ROOT / "data" / "deliveries" / "2026-03-14-baseline-v1"
OUTPUT_FILE = OUTPUT_DIR / "p0" / "compliance_rule.csv"
MANIFEST_FILE = OUTPUT_DIR / "reports" / "data_manifest.csv"
QUALITY_FILE = OUTPUT_DIR / "reports" / "compliance_rule_summary.md"
README_FILE = OUTPUT_DIR / "README.md"


def normalize_text(value: object) -> str:
    if value is None:
        return ""
    text = str(value).replace("　", " ").strip()
    text = re.sub(r"\s+", " ", text)
    return text


def normalize_multiline(value: object) -> str:
    text = normalize_text(value)
    text = text.replace("\n", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize_date(text: str) -> str:
    text = normalize_text(text)
    if not text:
        return ""
    match = re.search(r"(20\d{2})-(\d{1,2})-(\d{1,2})", text)
    if match:
        year, month, day = match.groups()
        return f"{year}-{int(month):02d}-{int(day):02d}"
    return ""


def title_from_path(path: Path) -> str:
    match = re.search(r"《([^》]+)》", path.name)
    if match:
        return match.group(1)
    return path.stem


def effective_date_from_path(path: Path) -> str:
    return normalize_date(path.name)


def make_rule_id(*parts: object) -> str:
    source = "|".join(normalize_text(part) for part in parts)
    return "rule_" + hashlib.sha1(source.encode("utf-8")).hexdigest()[:16]


def base_record(path: Path, source_sheet: str, rule_type: str, rule_level: str, source_title: str | None = None) -> dict[str, str]:
    return {
        "rule_id": "",
        "jurisdiction": "CN",
        "source_category": "regulation",
        "source_title": source_title or title_from_path(path),
        "source_file": path.relative_to(REPO_ROOT).as_posix(),
        "source_sheet": source_sheet,
        "effective_date": effective_date_from_path(path),
        "rule_type": rule_type,
        "rule_level": rule_level,
        "chapter": "",
        "article_no": "",
        "entity_name_cn": "",
        "entity_name_en": "",
        "inci_name": "",
        "cas_no": "",
        "applicable_scope": "",
        "limit_value": "",
        "warning_text": "",
        "requirement_text": "",
        "remarks": "",
    }


def read_docx_lines(path: Path) -> list[str]:
    doc = Document(path)
    lines: list[str] = []
    for paragraph in doc.paragraphs:
        text = normalize_multiline(paragraph.text)
        if text:
            lines.append(text)
    return lines


def read_doc_lines(path: Path) -> list[str]:
    result = subprocess.run(["textutil", "-convert", "txt", "-stdout", str(path)], capture_output=True, text=True, check=True)
    return [normalize_multiline(line) for line in result.stdout.splitlines() if normalize_multiline(line)]


def read_pdf_overview(path: Path) -> list[str]:
    reader = PdfReader(str(path))
    lines: list[str] = []
    for page in reader.pages[:8]:
        text = page.extract_text() or ""
        lines.extend(normalize_multiline(line) for line in text.splitlines() if normalize_multiline(line))
    return lines


def parse_clause_lines(path: Path, lines: list[str], rule_type: str) -> list[dict[str, str]]:
    title = title_from_path(path)
    chapter = ""
    article_no = ""
    article_text: list[str] = []
    records: list[dict[str, str]] = []
    title_candidates = {title, f"《{title}》", "附件", "附件1", "附件2", "中国食品药品检定研究院"}
    date_pattern = re.compile(r"^20\d{2}年\d{1,2}月\d{1,2}日$")
    chapter_pattern = re.compile(r"^第[一二三四五六七八九十百零]+章")
    article_patterns = [
        re.compile(r"^(第[一二三四五六七八九十百零]+条)"),
        re.compile(r"^([一二三四五六七八九十]+、)"),
        re.compile(r"^(\d+\.)"),
    ]

    def flush() -> None:
        nonlocal article_no, article_text
        text = normalize_multiline(" ".join(article_text))
        if not text:
            article_no = ""
            article_text = []
            return
        record = base_record(path, "", rule_type, "clause", source_title=title)
        record["chapter"] = chapter
        record["article_no"] = article_no or "preamble"
        record["requirement_text"] = text
        record["rule_id"] = make_rule_id(record["source_file"], record["chapter"], record["article_no"], text)
        records.append(record)
        article_no = ""
        article_text = []

    for raw_line in lines:
        line = normalize_multiline(raw_line)
        if not line or line in title_candidates or date_pattern.match(line):
            continue
        if chapter_pattern.match(line):
            flush()
            chapter = line
            continue
        matched = None
        for pattern in article_patterns:
            current = pattern.match(line)
            if current:
                matched = current.group(1)
                break
        if matched:
            flush()
            article_no = matched
            article_text = [line]
            continue
        if not article_text:
            article_text = [line]
        else:
            article_text.append(line)
    flush()
    return records


def extract_used_directory_docx(path: Path) -> list[dict[str, str]]:
    doc = Document(path)
    records: list[dict[str, str]] = []
    for table in doc.tables:
        header_found = False
        for row in table.rows:
            values = [normalize_multiline(cell.text) for cell in row.cells]
            if not any(values):
                continue
            if not header_found and values[0] == "序号":
                header_found = True
                continue
            if not header_found:
                continue
            if not values[0].isdigit():
                continue
            record = base_record(path, "table", "used_ingredient_directory_part2", "row")
            record["entity_name_cn"] = values[1] if len(values) > 1 else ""
            record["inci_name"] = values[2] if len(values) > 2 else ""
            record["entity_name_en"] = values[3] if len(values) > 3 else ""
            record["cas_no"] = values[4] if len(values) > 4 else ""
            record["limit_value"] = values[8] if len(values) > 8 else ""
            record["remarks"] = " | ".join(item for item in values[5:] if item)
            record["requirement_text"] = "已使用化妆品原料目录Ⅱ收录"
            record["rule_id"] = make_rule_id(record["source_file"], values[0], record["entity_name_cn"], record["inci_name"])
            records.append(record)
    return records


def parse_allowed_ingredients(path: Path) -> list[dict[str, str]]:
    wb = load_workbook(path, read_only=True, data_only=True)
    records: list[dict[str, str]] = []

    for sheet in wb.sheetnames:
        ws = wb[sheet]

        if "着色剂" in sheet:
            scope_labels = [normalize_multiline(value) for value in next(ws.iter_rows(min_row=4, max_row=4, values_only=True))[5:9]]
            for row in ws.iter_rows(min_row=5, values_only=True):
                values = [normalize_multiline(value) for value in row[:12]]
                if not any(values[1:5]):
                    continue
                scopes = [label for label, marker in zip(scope_labels, values[5:9]) if normalize_multiline(marker) == "＋"]
                record = base_record(path, sheet, "allowed_colorant", "row", source_title="准用化妆品成分")
                record["entity_name_cn"] = values[4]
                record["entity_name_en"] = values[2]
                record["inci_name"] = values[1]
                record["applicable_scope"] = " | ".join(scope for scope in scopes if scope)
                record["remarks"] = " | ".join(item for item in [f"颜色:{values[3]}" if values[3] else "", values[9], values[10], values[11]] if item)
                record["requirement_text"] = "准用着色剂目录收录"
                record["rule_id"] = make_rule_id(record["source_file"], sheet, values[0], record["inci_name"], record["entity_name_cn"])
                records.append(record)
            continue

        if "染发" in sheet:
            for row in ws.iter_rows(min_row=5, values_only=True):
                values = [normalize_multiline(value) for value in row[:10]]
                if not any(values[1:3]):
                    continue
                limits = []
                if values[3]:
                    limits.append(f"氧化型:{values[3]}")
                if values[4]:
                    limits.append(f"非氧化型:{values[4]}")
                record = base_record(path, sheet, "allowed_hair_dye", "row", source_title="准用化妆品成分")
                record["entity_name_cn"] = values[1]
                record["inci_name"] = values[2]
                record["limit_value"] = " | ".join(limits)
                record["applicable_scope"] = values[5]
                record["warning_text"] = values[6]
                record["remarks"] = " | ".join(item for item in [values[7], values[8], values[9]] if item)
                record["requirement_text"] = "准用染发剂目录收录"
                record["rule_id"] = make_rule_id(record["source_file"], sheet, values[0], record["entity_name_cn"], record["inci_name"])
                records.append(record)
            continue

        if "防晒" in sheet:
            for row in ws.iter_rows(min_row=4, values_only=True):
                values = [normalize_multiline(value) for value in row[:9]]
                if not any(values[1:4]):
                    continue
                record = base_record(path, sheet, "allowed_sunscreen", "row", source_title="准用化妆品成分")
                record["entity_name_cn"] = values[1]
                record["entity_name_en"] = values[2]
                record["inci_name"] = values[3]
                record["limit_value"] = values[4]
                record["applicable_scope"] = values[5]
                record["warning_text"] = values[6]
                record["remarks"] = values[7]
                record["requirement_text"] = "准用防晒剂目录收录"
                record["rule_id"] = make_rule_id(record["source_file"], sheet, values[0], record["entity_name_cn"], record["inci_name"], record["limit_value"])
                records.append(record)
            continue

        for row in ws.iter_rows(min_row=5, values_only=True):
            values = [normalize_multiline(value) for value in row[:11]]
            if not any(values[1:4]):
                continue
            record = base_record(path, sheet, "allowed_preservative", "row", source_title="准用化妆品成分")
            record["entity_name_cn"] = values[1]
            record["entity_name_en"] = values[2]
            record["inci_name"] = values[3]
            record["limit_value"] = values[4]
            record["applicable_scope"] = values[5]
            record["warning_text"] = values[6]
            record["remarks"] = " | ".join(item for item in values[7:11] if item)
            record["requirement_text"] = "准用防腐剂目录收录"
            record["rule_id"] = make_rule_id(record["source_file"], sheet, values[0], record["entity_name_cn"], record["inci_name"])
            records.append(record)

    wb.close()
    return records


def parse_banned_ingredients(path: Path) -> list[dict[str, str]]:
    ws = load_workbook(path, read_only=True, data_only=True)["禁用组分2026.1"]
    records: list[dict[str, str]] = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        values = [normalize_multiline(value) for value in row[:10]]
        if not any(values[2:4]):
            continue
        record = base_record(path, "禁用组分2026.1", "prohibited_ingredient", "row", source_title="禁用成分目录")
        record["entity_name_cn"] = values[2]
        record["entity_name_en"] = values[3]
        record["cas_no"] = values[8]
        record["remarks"] = " | ".join(item for item in [values[4], values[5], values[6], values[7], values[9]] if item)
        record["requirement_text"] = "禁用成分，不得用于化妆品"
        record["rule_id"] = make_rule_id(record["source_file"], values[1], record["entity_name_cn"], record["cas_no"])
        records.append(record)
    return records


def parse_restricted_ingredients(path: Path) -> list[dict[str, str]]:
    ws = load_workbook(path, read_only=True, data_only=True)["限用组分2026.1"]
    records: list[dict[str, str]] = []
    for row in ws.iter_rows(min_row=5, values_only=True):
        values = [normalize_multiline(value) for value in row[:12]]
        if not any(values[1:4]):
            continue
        record = base_record(path, "限用组分2026.1", "restricted_ingredient", "row", source_title="限用组分目录")
        record["entity_name_cn"] = values[1]
        record["entity_name_en"] = values[2]
        record["inci_name"] = values[3]
        record["applicable_scope"] = values[4]
        record["limit_value"] = values[5]
        record["warning_text"] = values[7]
        record["remarks"] = " | ".join(item for item in [values[6], values[8], values[9], values[10], values[11]] if item)
        record["requirement_text"] = "限用组分，须满足适用范围和浓度限制"
        record["rule_id"] = make_rule_id(record["source_file"], values[0], record["entity_name_cn"], record["applicable_scope"], record["limit_value"])
        records.append(record)
    return records


def parse_banned_plant_ingredients(path: Path) -> list[dict[str, str]]:
    ws = load_workbook(path, read_only=True, data_only=True)["（中国）化妆品禁用植（动）物原料目录26.1"]
    records: list[dict[str, str]] = []
    for row in ws.iter_rows(min_row=3, values_only=True):
        values = [normalize_multiline(value) for value in row[:8]]
        if not any(values[1:4]):
            continue
        record = base_record(path, "（中国）化妆品禁用植（动）物原料目录26.1", "prohibited_plant_animal_ingredient", "row", source_title="禁用植（动）物原料目录")
        record["entity_name_cn"] = values[1]
        record["entity_name_en"] = values[2]
        record["cas_no"] = values[4]
        record["remarks"] = " | ".join(item for item in [values[3], values[5], values[6], values[7]] if item)
        record["requirement_text"] = "禁用植（动）物原料，不得用于化妆品"
        record["rule_id"] = make_rule_id(record["source_file"], values[0], record["entity_name_cn"], record["cas_no"])
        records.append(record)
    return records


def parse_used_ingredients_xlsx(path: Path) -> list[dict[str, str]]:
    ws = load_workbook(path, read_only=True, data_only=True)["源数据 (3)"]
    records: list[dict[str, str]] = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        values = [normalize_multiline(value) for value in row[:4]]
        if not any(values[1:3]):
            continue
        record = base_record(path, "源数据 (3)", "used_ingredient_directory", "row", source_title="已使用化妆品原料目录Ⅰ")
        record["entity_name_cn"] = values[1]
        record["inci_name"] = values[2]
        record["remarks"] = values[3]
        record["requirement_text"] = "已使用化妆品原料目录Ⅰ收录"
        record["rule_id"] = make_rule_id(record["source_file"], values[0], record["entity_name_cn"], record["inci_name"])
        records.append(record)
    return records


def parse_usage_reference_xlsx(path: Path, sheet_name: str, rule_type: str, title: str) -> list[dict[str, str]]:
    ws = load_workbook(path, read_only=True, data_only=True)[sheet_name]
    records: list[dict[str, str]] = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        values = [normalize_multiline(value) for value in row[:10]]
        if not any(values[2:4]):
            continue
        record = base_record(path, sheet_name, rule_type, "row", source_title=title)
        record["entity_name_cn"] = values[2]
        record["inci_name"] = values[3]
        record["applicable_scope"] = " | ".join(item for item in [values[4], values[5]] if item)
        record["limit_value"] = values[6]
        record["remarks"] = " | ".join(item for item in [values[1], values[7], values[8], values[9]] if item)
        record["requirement_text"] = title
        record["rule_id"] = make_rule_id(record["source_file"], values[0], record["entity_name_cn"], record["applicable_scope"], record["limit_value"])
        records.append(record)
    return records


def parse_pdf_summary(path: Path) -> list[dict[str, str]]:
    lines = read_pdf_overview(path)
    summary_lines = []
    for line in lines:
        if len(summary_lines) >= 20:
            break
        summary_lines.append(line)
    text = normalize_multiline(" ".join(summary_lines))
    if not text:
        return []
    record = base_record(path, "", "technical_standard_overview", "summary")
    record["requirement_text"] = text[:4000]
    record["rule_id"] = make_rule_id(record["source_file"], record["source_title"], record["requirement_text"][:120])
    return [record]


def build_all_records() -> pd.DataFrame:
    root = RAW_DIR
    records: list[dict[str, str]] = []
    ingredient_dir = root / "原料目录"
    records.extend(parse_allowed_ingredients(ingredient_dir / "(中国)准用化妆品成分26.1.xlsx"))
    records.extend(parse_banned_ingredients(ingredient_dir / "(中国)禁用成分26.1.xlsx"))
    records.extend(parse_restricted_ingredients(ingredient_dir / "(中国)限用组分26.1.xlsx"))
    records.extend(parse_used_ingredients_xlsx(ingredient_dir / "《已使用化妆品原料目录》Ⅰ.xlsx"))
    records.extend(extract_used_directory_docx(ingredient_dir / "《已使用化妆品原料目录》Ⅱ.docx"))
    records.extend(parse_usage_reference_xlsx(ingredient_dir / "《国际化妆品安全评估数据索引》收录的部分原料使用信息.xlsx", "数据处理结果", "safety_assessment_reference", "国际化妆品安全评估数据索引收录原料使用信息"))
    records.extend(parse_usage_reference_xlsx(ingredient_dir / "已上市产品原料使用信息.xlsx", "处理结果", "marketed_usage_reference", "已上市产品原料使用信息"))
    records.extend(parse_banned_plant_ingredients(ingredient_dir / "（中国）化妆品禁用植（动）物原料目录26.1.xlsx"))

    regulation_dir = root / "法规文件"
    for path in sorted(regulation_dir.iterdir()):
        suffix = path.suffix.lower()
        title = title_from_path(path)
        if suffix == ".docx":
            records.extend(parse_clause_lines(path, read_docx_lines(path), rule_type=title))
        elif suffix == ".doc":
            records.extend(parse_clause_lines(path, read_doc_lines(path), rule_type=title))
        elif suffix == ".pdf":
            records.extend(parse_pdf_summary(path))

    frame = pd.DataFrame(records)
    frame = frame.fillna("")
    frame = frame.drop_duplicates(subset=["rule_id"]).sort_values(by=["rule_type", "source_title", "source_sheet", "article_no", "entity_name_cn"]).reset_index(drop=True)
    return frame


def write_outputs(frame: pd.DataFrame) -> None:
    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    frame.to_csv(OUTPUT_FILE, index=False, encoding="utf-8-sig")

    manifest = pd.read_csv(MANIFEST_FILE)
    manifest = manifest[manifest["table_name"] != "compliance_rule"]
    manifest = pd.concat([
        manifest,
        pd.DataFrame([
            {
                "table_name": "compliance_rule",
                "level": "P0",
                "rows": len(frame),
                "path": "p0/compliance_rule.csv",
            }
        ]),
    ], ignore_index=True)
    manifest.to_csv(MANIFEST_FILE, index=False, encoding="utf-8-sig")

    counts = frame.groupby("rule_type").size().sort_values(ascending=False)
    lines = ["# Compliance Rule Summary", "", f"- rows: {len(frame)}", f"- source files covered: {frame['source_file'].nunique()}", "", "## Count By Rule Type", ""]
    for rule_type, count in counts.items():
        lines.append(f"- {rule_type}: {int(count)}")
    QUALITY_FILE.write_text("\n".join(lines), encoding="utf-8")

    readme = README_FILE.read_text(encoding="utf-8")
    if "compliance_rule.csv" not in readme:
        readme = readme.replace(
            "- `p1/`: restricted package containing raw review text\n",
            "- `p1/`: restricted package containing raw review text\n- `p0/compliance_rule.csv`: first-pass structured compliance rules\n",
        )
        README_FILE.write_text(readme, encoding="utf-8")


def main() -> None:
    frame = build_all_records()
    write_outputs(frame)


if __name__ == "__main__":
    main()
